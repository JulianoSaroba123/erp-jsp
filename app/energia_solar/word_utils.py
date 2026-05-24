"""
Utilitário para processar templates Word (.docx) com substituição de variáveis
"""
from docx import Document
from docx.shared import Inches
import io
import os
import zipfile
import tempfile
import logging
import re
from xml.sax.saxutils import escape


def _to_float(value, default=0.0):
    try:
        return float(value)
    except Exception:
        return default


def _bar(valor, maximo, largura=16, char='#'):
    """Gera barra ASCII proporcional para visual tipo dashboard."""
    maximo = max(1.0, _to_float(maximo, 1.0))
    valor = max(0.0, _to_float(valor, 0.0))
    preenchido = int(round((valor / maximo) * largura))
    preenchido = max(0, min(largura, preenchido))
    return (char * preenchido) + ('.' * (largura - preenchido))


def _to_float_text(value, default=0.0):
    """Converte texto numérico em float, aceitando vírgula decimal."""
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    try:
        txt = str(value).strip().replace('R$', '').replace(' ', '')
        if ',' in txt and '.' in txt:
            txt = txt.replace('.', '').replace(',', '.')
        else:
            txt = txt.replace(',', '.')
        return float(txt)
    except Exception:
        return default


def _extract_first_float(value, default=0.0):
    """Extrai o primeiro número de uma string (ex.: '3,0 anos' -> 3.0)."""
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    txt = str(value)
    m = re.search(r'-?\d+[\.,]?\d*', txt)
    if not m:
        return default
    try:
        return float(m.group(0).replace(',', '.'))
    except Exception:
        return default


def _fmt_kwh(value):
    """Formata energia em kWh com separador de milhar estilo PT-BR."""
    n = int(round(_to_float_text(value, 0)))
    return f'{n:,}'.replace(',', '.')


def _fmt_rs(value):
    """Formata moeda em real com 2 casas no padrão PT-BR."""
    n = _to_float_text(value, 0.0)
    txt = f'{n:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
    return f'R$ {txt}'


def _normalizar_tipo_instalacao(valor):
    txt = (str(valor or '').strip().lower())
    if 'tri' in txt:
        return 'trifasica'
    if 'bi' in txt:
        return 'bifasica'
    return 'monofasica'


def _consumo_minimo_kwh_por_tipo(tipo_instalacao):
    tipo = _normalizar_tipo_instalacao(tipo_instalacao)
    mapa = {
        'monofasica': 30.0,
        'bifasica': 50.0,
        'trifasica': 100.0,
    }
    return mapa.get(tipo, 30.0)


def _normalizar_variaveis_financeiras(variaveis):
    """Garante campos financeiros essenciais para evitar placeholders zerados."""
    consumo = _to_float_text(variaveis.get('consumo_kwh_mes'), 0)

    tarifa = _to_float_text(variaveis.get('tarifa_kwh'), 0)
    if tarifa <= 0:
        tarifa = _to_float_text(variaveis.get('tarifa_energia'), 0)
    if tarifa <= 0:
        tarifa = _to_float_text(variaveis.get('preco_kwh'), 0)

    fatura_sem = _to_float_text(variaveis.get('fatura_sem_sistema'), 0)
    if fatura_sem <= 0:
        fatura_sem = _to_float_text(variaveis.get('fatura_media_mensal_sem_sistema'), 0)
    if fatura_sem <= 0:
        fatura_sem = _to_float_text(variaveis.get('conta_luz_atual'), 0)

    if fatura_sem <= 0 and consumo > 0 and tarifa > 0:
        fatura_sem = consumo * tarifa

    if tarifa <= 0 and consumo > 0 and fatura_sem > 0:
        tarifa = fatura_sem / max(consumo, 1.0)

    if tarifa <= 0:
        tarifa = 0.85

    tipo_instalacao = variaveis.get('tipo_instalacao') or variaveis.get('TIPO_INSTALACAO') or 'monofasica'
    consumo_minimo_kwh = _consumo_minimo_kwh_por_tipo(tipo_instalacao)

    adicionais = 0.0
    for chave in [
        'iluminacao_publica',
        'demais_custos',
        'outros_custos',
        'adicionais_fatura',
        'cip',
        'taxa_disponibilidade',
        'taxa_iluminacao_publica',
    ]:
        adicionais += _to_float_text(variaveis.get(chave), 0)

    fatura_minima_tecnica = (consumo_minimo_kwh * tarifa) + adicionais

    economia_mensal = _to_float_text(variaveis.get('economia_mensal'), 0)
    if economia_mensal <= 0:
        economia_mensal = _to_float_text(variaveis.get('ECONOMIA_MENSAL'), 0)
    if economia_mensal <= 0 and fatura_sem > 0:
        economia_mensal = fatura_sem * 0.9

    fatura_com = _to_float_text(variaveis.get('fatura_com_sistema'), 0)
    if fatura_com <= 0:
        fatura_com = _to_float_text(variaveis.get('fatura_minima'), 0)
    if fatura_com <= 0:
        fatura_com = _to_float_text(variaveis.get('conta_luz_futura'), 0)

    if fatura_com <= 0 and fatura_sem > 0:
        fatura_com = max(fatura_sem - economia_mensal, fatura_minima_tecnica)

    if fatura_com > 0:
        fatura_com = max(fatura_com, fatura_minima_tecnica)

    if fatura_sem > 0 and fatura_sem < fatura_minima_tecnica:
        fatura_sem = fatura_minima_tecnica

    if fatura_sem <= 0 and fatura_com > 0:
        fatura_sem = fatura_com / 0.1

    if fatura_com > fatura_sem > 0:
        fatura_com = fatura_sem

    acrescimo_anual = _to_float_text(
        variaveis.get('acrescimo_anual_percentual')
        or variaveis.get('reajuste_anual_energia')
        or variaveis.get('reajuste_anual')
        or 10.0,
        10.0,
    )

    variaveis.update({
        'tarifa_kwh': f'{tarifa:.4f}',
        'preco_kwh': f'R$ {tarifa:.4f}',
        'tarifa_energia': f'R$ {tarifa:.4f}',
        'fatura_media_mensal_sem_sistema': f'{fatura_sem:.2f}',
        'fatura_sem_sistema': f'{fatura_sem:.2f}',
        'fatura_com_sistema': f'{fatura_com:.2f}',
        'fatura_minima': f'{fatura_com:.2f}',
        'fatura_minima_tecnica': f'{fatura_minima_tecnica:.2f}',
        'conta_luz_atual': f'{fatura_sem:.2f}',
        'conta_luz_futura': f'{fatura_com:.2f}',
        'consumo_minimo_kwh': f'{consumo_minimo_kwh:.0f}',
        'acrescimo_anual_percentual': f'{acrescimo_anual:.2f}',
        'fatura_sem_sistema_rs': _fmt_rs(fatura_sem),
        'fatura_com_sistema_rs': _fmt_rs(fatura_com),
        'conta_luz_atual_rs': _fmt_rs(fatura_sem),
        'conta_luz_futura_rs': _fmt_rs(fatura_com),
    })

    return variaveis


def _normalizar_series_12m(consumo_base, geracao_base):
    """Cria perfil mensal mais orgânico para consumo e geração."""
    # Consumo tende a oscilar menos que geração, simulando sazonalidade e uso.
    fatores_consumo = [1.03, 1.02, 1.01, 1.00, 0.98, 0.97, 0.96, 0.97, 0.99, 1.01, 1.03, 1.04]
    fatores_geracao = [0.93, 0.92, 0.95, 0.98, 1.02, 1.06, 1.08, 1.07, 1.03, 1.00, 0.97, 0.94]
    consumo = [consumo_base * f for f in fatores_consumo]
    geracao = [geracao_base * f for f in fatores_geracao]
    return consumo, geracao


def _serie_consumo_12_meses_variaveis(variaveis):
    meses = ['jan', 'fev', 'mar', 'abr', 'mai', 'jun', 'jul', 'ago', 'set', 'out', 'nov', 'dez']
    base = _to_float_text(variaveis.get('consumo_kwh_mes', 0), 0)

    valores = []
    for m in meses:
        v = (
            variaveis.get(f'consumo_{m}')
            or variaveis.get(f'consumo_kwh_{m}')
            or variaveis.get(f'consumo_{m}_kwh')
        )
        valores.append(_to_float_text(v, 0))

    if any(v > 0 for v in valores):
        return [v if v > 0 else base for v in valores]

    consumo, _ = _normalizar_series_12m(base, _to_float_text(variaveis.get('geracao_estimada_mes', 0), 0))
    return consumo


def _serie_geracao_12_meses_variaveis(variaveis):
    meses = ['jan', 'fev', 'mar', 'abr', 'mai', 'jun', 'jul', 'ago', 'set', 'out', 'nov', 'dez']
    base = _to_float_text(variaveis.get('geracao_estimada_mes', 0), 0)

    valores = []
    for m in meses:
        v = (
            variaveis.get(f'geracao_{m}')
            or variaveis.get(f'geracao_kwh_{m}')
            or variaveis.get(f'geracao_{m}_kwh')
        )
        valores.append(_to_float_text(v, 0))

    if any(v > 0 for v in valores):
        return [v if v > 0 else base for v in valores]

    _, geracao = _normalizar_series_12m(_to_float_text(variaveis.get('consumo_kwh_mes', 0), 0), base)
    return geracao


def _placeholder_variantes(chave):
    c = str(chave)
    return {
        f'[{c}]', f'[{c.upper()}]', f'[{c.lower()}]',
        f'{{{{{c}}}}}', f'{{{{{c.upper()}}}}}', f'{{{{{c.lower()}}}}}',
        f'{{{c}}}', f'{{{c.upper()}}}', f'{{{c.lower()}}}',
        f'<<{c}>>', f'<<{c.upper()}>>', f'<<{c.lower()}>>',
        f'${{{c}}}', f'${{{c.upper()}}}', f'${{{c.lower()}}}',
        f'%{c}%', f'%{c.upper()}%', f'%{c.lower()}%',
    }


def _calcular_tarifa_financeira_projeto(projeto, tarifa_padrao=0.85):
    """Obtém tarifa efetiva considerando concessionária e aplicações TE/TUSD."""
    tarifa = _extract_first_float(getattr(projeto, 'tarifa_kwh', None), 0)

    try:
        from app.concessionaria.concessionaria_model import Concessionaria

        concessionaria_id = getattr(projeto, 'concessionaria_id', None)
        if concessionaria_id:
            conc = Concessionaria.query.get(concessionaria_id)
            if conc:
                te = _to_float(getattr(conc, 'te', 0), 0)
                tusd = _to_float(getattr(conc, 'tusd', 0), 0)
                pis = _to_float(getattr(conc, 'pis', 0), 0)
                cofins = _to_float(getattr(conc, 'cofins', 0), 0)
                icms = _to_float(getattr(conc, 'icms', 0), 0)

                aplicar_pis_te = getattr(projeto, 'aplicar_pis_te', True)
                aplicar_cofins_te = getattr(projeto, 'aplicar_cofins_te', True)
                aplicar_icms_te = getattr(projeto, 'aplicar_icms_te', True)
                aplicar_pis_tusd = getattr(projeto, 'aplicar_pis_tusd', True)
                aplicar_cofins_tusd = getattr(projeto, 'aplicar_cofins_tusd', True)
                aplicar_icms_tusd = getattr(projeto, 'aplicar_icms_tusd', True)

                perc_te = (
                    (pis if aplicar_pis_te else 0)
                    + (cofins if aplicar_cofins_te else 0)
                    + (icms if aplicar_icms_te else 0)
                )
                perc_tusd = (
                    (pis if aplicar_pis_tusd else 0)
                    + (cofins if aplicar_cofins_tusd else 0)
                    + (icms if aplicar_icms_tusd else 0)
                )

                tarifa_calc = (te * (1 + perc_te / 100.0)) + (tusd * (1 + perc_tusd / 100.0))
                if tarifa_calc > 0:
                    tarifa = tarifa_calc
    except Exception:
        pass

    if tarifa <= 0:
        tarifa = tarifa_padrao
    return tarifa


def _gerar_imagem_grafico_consumo_geracao_12m(variaveis):
    """Gera gráfico visual (barra + linha) para consumo x geração em 12 meses."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        matplotlib.set_loglevel('warning')
        logging.getLogger('matplotlib').setLevel(logging.WARNING)
        import matplotlib.pyplot as plt
        import matplotlib.patheffects as pe
    except Exception:
        return None

    meses = ['Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']
    consumo = _serie_consumo_12_meses_variaveis(variaveis)
    geracao = _serie_geracao_12_meses_variaveis(variaveis)

    simult = _to_float_text(variaveis.get('simultaneidade', 35), 35)
    consumo_simult = [c * (simult / 100.0) for c in consumo]

    # Valores zerados indicam que não há base suficiente para desenhar o gráfico.
    if max(consumo + geracao) <= 0:
        return None

    fig, ax = plt.subplots(figsize=(11.2, 4.9), dpi=190)
    fig.patch.set_facecolor('#FFFFFF')
    ax.set_facecolor('#F7FBFF')
    x = list(range(len(meses)))
    largura = 0.36

    consumo_bar = ax.bar(
        [i - largura / 2 for i in x],
        consumo,
        width=largura,
        color='#27B4FF',
        edgecolor='#0A7FC2',
        linewidth=0.9,
        label='Consumo (kWh)'
    )
    geracao_bar = ax.bar(
        [i + largura / 2 for i in x],
        geracao,
        width=largura,
        color='#6FF2A9',
        edgecolor='#2EA866',
        linewidth=0.9,
        label='Geracao (kWh)'
    )
    line_simult, = ax.plot(x, consumo_simult, color='#FF5E84', marker='o', markersize=4.3, linewidth=2.0, label='Consumo simultaneo')

    # Pseudo-3D: sombra nas barras e glow na linha.
    for b in list(consumo_bar) + list(geracao_bar):
        b.set_path_effects([pe.SimplePatchShadow(offset=(1.2, -1.2), alpha=0.28), pe.Normal()])
    line_simult.set_path_effects([pe.Stroke(linewidth=3.4, foreground='#FFC4D2', alpha=0.8), pe.Normal()])

    ax.set_xticks(x)
    ax.set_xticklabels(meses, fontsize=10, fontweight='bold')
    ax.tick_params(axis='y', labelsize=10)
    ax.set_ylabel('kWh/mes', fontsize=10, color='#24384D', fontweight='bold')
    ax.set_ylim(0, max(max(consumo), max(geracao)) * 1.22)
    ax.grid(axis='y', color='#DFE7EF', linewidth=0.8, alpha=0.9)
    ax.set_axisbelow(True)

    for side in ['top', 'right']:
        ax.spines[side].set_visible(False)
    ax.spines['left'].set_color('#C9D6E2')
    ax.spines['bottom'].set_color('#C9D6E2')

    # Destaque do ganho anual aproximado para leitura executiva.
    energia_gerada_ano = sum(geracao)
    energia_consumida_ano = sum(consumo)
    delta_ano = energia_gerada_ano - energia_consumida_ano
    ax.text(
        0.01,
        0.98,
        f'Geracao anual: {_fmt_kwh(energia_gerada_ano)} kWh  |  Consumo anual: {_fmt_kwh(energia_consumida_ano)} kWh  |  Saldo: {_fmt_kwh(delta_ano)} kWh',
        transform=ax.transAxes,
        ha='left',
        va='top',
        fontsize=9,
        color='#1F3A56',
        bbox=dict(boxstyle='round,pad=0.25', fc='#EEF5FF', ec='#C9DDF7', lw=0.7)
    )

    # Rótulos discretos para melhorar leitura sem poluir.
    for idx in [0, 5, 11]:
        ax.text(consumo_bar[idx].get_x() + consumo_bar[idx].get_width() / 2, consumo[idx] + 8, _fmt_kwh(consumo[idx]),
                ha='center', va='bottom', fontsize=8, color='#0A5A87', fontweight='bold')
        ax.text(geracao_bar[idx].get_x() + geracao_bar[idx].get_width() / 2, geracao[idx] + 8, _fmt_kwh(geracao[idx]),
                ha='center', va='bottom', fontsize=8, color='#1C7547', fontweight='bold')

    ax.set_title('Geracao de energia - estimativa mes a mes', loc='left', fontsize=14, fontweight='bold', color='#0E2C4D', pad=12)
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.12), ncol=3, fontsize=9, frameon=False)

    fig.tight_layout()

    out = io.BytesIO()
    fig.savefig(out, format='png', bbox_inches='tight', transparent=False)
    plt.close(fig)
    out.seek(0)
    return out


def _gerar_imagem_tabela_12m(variaveis):
    """Gera tabela visual de 12 meses para inserir no DOCX."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        matplotlib.set_loglevel('warning')
        logging.getLogger('matplotlib').setLevel(logging.WARNING)
        import matplotlib.pyplot as plt
    except Exception:
        return None

    meses = ['Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']
    consumo = _serie_consumo_12_meses_variaveis(variaveis)
    geracao = _serie_geracao_12_meses_variaveis(variaveis)
    tarifa = _to_float_text(variaveis.get('tarifa_kwh', 0), 0)
    if tarifa <= 0:
        tarifa = _to_float_text(variaveis.get('tarifa_energia', 0), 0)
    if tarifa <= 0:
        tarifa = _to_float_text(variaveis.get('preco_kwh', 0), 0)

    fatura_sem_ref = _to_float_text(variaveis.get('fatura_sem_sistema', 0), 0)
    if fatura_sem_ref <= 0:
        fatura_sem_ref = _to_float_text(variaveis.get('fatura_media_mensal_sem_sistema', 0), 0)

    fatura_min = _to_float_text(variaveis.get('fatura_minima', 0), 0)
    if fatura_min <= 0:
        fatura_min = _to_float_text(variaveis.get('fatura_com_sistema', 0), 0)

    media_consumo = sum(consumo) / max(len(consumo), 1)
    if tarifa <= 0 and media_consumo > 0 and fatura_sem_ref > 0:
        tarifa = fatura_sem_ref / media_consumo
    if tarifa <= 0:
        tarifa = 0.85

    if fatura_min <= 0 and fatura_sem_ref > 0:
        fatura_min = fatura_sem_ref * 0.1

    if max(consumo + geracao) <= 0:
        return None

    linhas = []
    total_consumo = 0.0
    total_geracao = 0.0
    total_sem = 0.0
    total_com = 0.0

    for i, mes in enumerate(meses):
        c = consumo[i]
        g = geracao[i]
        saldo = g - c
        sem_sistema = c * tarifa
        com_sistema = fatura_min

        total_consumo += c
        total_geracao += g
        total_sem += sem_sistema
        total_com += com_sistema

        linhas.append([mes, _fmt_kwh(c), _fmt_kwh(g), _fmt_kwh(saldo), _fmt_rs(tarifa), _fmt_rs(sem_sistema), _fmt_rs(com_sistema)])

    linhas.append(['Total 12m', _fmt_kwh(total_consumo), _fmt_kwh(total_geracao), _fmt_kwh(total_geracao - total_consumo), '-', _fmt_rs(total_sem), _fmt_rs(total_com)])

    colunas = ['Mes', 'Consumo (kWh)', 'Geracao (kWh)', 'Saldo (kWh)', 'Tarifa', 'Fatura sem sistema', 'Fatura com sistema']

    fig_h = 0.46 * (len(linhas) + 1)
    fig, ax = plt.subplots(figsize=(10.8, fig_h), dpi=170)
    fig.patch.set_facecolor('#FFFFFF')
    ax.axis('off')

    tabela = ax.table(cellText=linhas, colLabels=colunas, loc='center', cellLoc='center')
    tabela.auto_set_font_size(False)
    tabela.set_fontsize(9)
    tabela.scale(1, 1.24)

    # Header com destaque visual
    for c_idx in range(len(colunas)):
        cell = tabela[(0, c_idx)]
        cell.set_facecolor('#0F2F53')
        cell.set_edgecolor('#28507A')
        cell.set_text_props(weight='bold', color='#E9F6FF')

    # Zebra striping para facilitar leitura de linha.
    for r_idx in range(1, len(linhas)):
        bg = '#FFFFFF' if r_idx % 2 == 1 else '#EEF6FF'
        for c_idx in range(len(colunas)):
            cell = tabela[(r_idx, c_idx)]
            cell.set_facecolor(bg)
            cell.set_edgecolor('#E3EAF2')

    # Linha de total em destaque suave
    total_row_index = len(linhas)
    for c_idx in range(len(colunas)):
        cell = tabela[(total_row_index, c_idx)]
        cell.set_facecolor('#DDFBEA')
        cell.set_edgecolor('#9FD4B5')
        cell.set_text_props(weight='bold', color='#0E5A2D')

    ax.set_title('Resumo financeiro e energetico - 12 meses', loc='left', fontsize=12, fontweight='bold', color='#0E2C4D', pad=10)

    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format='png', bbox_inches='tight', transparent=False)
    plt.close(fig)
    out.seek(0)
    return out


def _serie_economia_25_anos_variaveis(variaveis):
    """Gera série de economia anual e acumulada para 25 anos."""
    investimento = _to_float_text(
        variaveis.get('valor_total')
        or variaveis.get('orcamento_valor_nota')
        or variaveis.get('VALOR_INVESTIMENTO')
        or 0,
        0.0,
    )

    economia_mensal = _to_float_text(
        variaveis.get('economia_mensal')
        or variaveis.get('ECONOMIA_MENSAL')
        or 0,
        0,
    )

    if economia_mensal <= 0:
        economia_anual = _to_float_text(
            variaveis.get('economia_anual')
            or variaveis.get('ECONOMIA_ANUAL')
            or 0,
            0,
        )
        if economia_anual > 0:
            economia_mensal = economia_anual / 12.0

    if economia_mensal <= 0:
        payback_anos = _extract_first_float(
            variaveis.get('payback_ano_lei_14300')
            or variaveis.get('payback_roi_lei_14300')
            or variaveis.get('PAYBACK')
            or variaveis.get('payback')
            or variaveis.get('payback_anos')
            or 0,
            0.0,
        )
        if investimento > 0 and payback_anos > 0:
            economia_mensal = investimento / (payback_anos * 12.0)

    if economia_mensal <= 0:
        fatura_sem = _to_float_text(variaveis.get('fatura_media_mensal_sem_sistema'), 0)
        fatura_com = _to_float_text(variaveis.get('fatura_minima'), 0)
        if fatura_sem > 0:
            economia_mensal = max(fatura_sem - fatura_com, fatura_sem * 0.85)

    if economia_mensal <= 0:
        consumo = _to_float_text(variaveis.get('consumo_kwh_mes'), 0)
        tarifa = _to_float_text(variaveis.get('tarifa_kwh'), 0)
        if consumo > 0 and tarifa > 0:
            economia_mensal = consumo * tarifa * 0.9

    if economia_mensal <= 0:
        consumo_serie = _serie_consumo_12_meses_variaveis(variaveis)
        consumo_medio = (sum(consumo_serie) / len(consumo_serie)) if consumo_serie else 0
        tarifa = _to_float_text(variaveis.get('tarifa_kwh'), 0)
        if tarifa <= 0:
            tarifa = _to_float_text(variaveis.get('preco_kwh'), 0)
        if tarifa <= 0:
            tarifa = 0.85

        if consumo_medio > 0 and tarifa > 0:
            tipo_instalacao = variaveis.get('tipo_instalacao') or variaveis.get('TIPO_INSTALACAO') or 'monofasica'
            consumo_minimo_kwh = _consumo_minimo_kwh_por_tipo(tipo_instalacao)
            adicionais = _to_float_text(variaveis.get('iluminacao_publica'), 0) + _to_float_text(variaveis.get('demais_custos'), 0)
            fatura_sem = consumo_medio * tarifa
            fatura_com = (consumo_minimo_kwh * tarifa) + adicionais
            economia_mensal = max(fatura_sem - fatura_com, fatura_sem * 0.6)

    reajuste = _to_float_text(
        variaveis.get('acrescimo_anual_percentual')
        or variaveis.get('reajuste_anual_energia')
        or variaveis.get('reajuste_anual')
        or 10.0,
        10.0,
    )
    taxa = reajuste / 100.0

    economia_base_anual = economia_mensal * 12
    anual = []
    acumulada = []
    soma = 0.0
    for ano in range(1, 26):
        v = economia_base_anual * ((1 + taxa) ** (ano - 1))
        soma += v
        anual.append(v)
        acumulada.append(soma)
    return anual, acumulada


def _gerar_imagem_grafico_economia_25_anos(variaveis):
    """Gera gráfico visual da projeção de economia para 25 anos."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        matplotlib.set_loglevel('warning')
        logging.getLogger('matplotlib').setLevel(logging.WARNING)
        import matplotlib.pyplot as plt
        import matplotlib.patheffects as pe
    except Exception:
        return None

    anual, acumulada = _serie_economia_25_anos_variaveis(variaveis)
    if max(acumulada) <= 0:
        return None

    anos = list(range(1, 26))

    fig, ax1 = plt.subplots(figsize=(11.2, 4.9), dpi=190)
    fig.patch.set_facecolor('#FFFFFF')
    ax1.set_facecolor('#F7FBFF')

    bars = ax1.bar(
        anos,
        anual,
        color='#41B8FF',
        edgecolor='#0E7FC3',
        linewidth=0.9,
        label='Economia anual (R$)'
    )

    ax2 = ax1.twinx()
    linha_acum, = ax2.plot(anos, acumulada, color='#12B76A', linewidth=2.4, marker='o', markersize=3.8, label='Economia acumulada (R$)')
    ax2.fill_between(anos, acumulada, color='#BFF2D7', alpha=0.38)

    for b in bars:
        b.set_path_effects([pe.SimplePatchShadow(offset=(1.2, -1.2), alpha=0.28), pe.Normal()])
    linha_acum.set_path_effects([pe.Stroke(linewidth=3.6, foreground='#9BF0C8', alpha=0.85), pe.Normal()])

    ax1.set_xlim(0.2, 25.8)
    ax1.set_xticks([1, 3, 5, 7, 10, 13, 16, 19, 22, 25])
    ax1.set_xlabel('Ano', fontsize=10, color='#24384D', fontweight='bold')
    ax1.set_ylabel('Economia anual (R$)', fontsize=10, color='#1C5D8F', fontweight='bold')
    ax2.set_ylabel('Economia acumulada (R$)', fontsize=10, color='#147A49', fontweight='bold')
    ax1.tick_params(axis='x', labelsize=10)
    ax1.tick_params(axis='y', labelsize=10)
    ax2.tick_params(axis='y', labelsize=10)

    ax1.grid(axis='y', color='#DFE7EF', linewidth=0.8, alpha=0.9)
    ax1.set_axisbelow(True)

    for side in ['top', 'right']:
        ax1.spines[side].set_visible(False)
    ax1.spines['left'].set_color('#C9D6E2')
    ax1.spines['bottom'].set_color('#C9D6E2')
    ax2.spines['top'].set_visible(False)
    ax2.spines['left'].set_visible(False)
    ax2.spines['right'].set_color('#C9D6E2')

    # Destaque executivo
    ax1.text(
        0.01,
        0.98,
        f'Acumulado em 25 anos: {_fmt_rs(acumulada[-1])}  |  Economia no 1o ano: {_fmt_rs(anual[0])}',
        transform=ax1.transAxes,
        ha='left',
        va='top',
        fontsize=9,
        color='#1F3A56',
        bbox=dict(boxstyle='round,pad=0.25', fc='#EEF5FF', ec='#C9DDF7', lw=0.7)
    )

    # Rótulos pontuais para leitura rápida
    for idx in [0, 9, 24]:
        b = bars[idx]
        ax1.text(
            b.get_x() + b.get_width() / 2,
            b.get_height() + (max(anual) * 0.015),
            _fmt_rs(anual[idx]).replace('R$ ', ''),
            ha='center',
            va='bottom',
            fontsize=8,
            color='#2F6AA6'
        )

    ax2.text(25, acumulada[-1], _fmt_rs(acumulada[-1]), fontsize=8, color='#147A49', ha='right', va='bottom', fontweight='bold')

    ax1.set_title('Estimativa de geracao para os proximos 25 anos', loc='left', fontsize=14, fontweight='bold', color='#0E2C4D', pad=12)

    h1, l1 = ax1.get_legend_handles_labels()
    h2, l2 = ax2.get_legend_handles_labels()
    ax1.legend(h1 + h2, l1 + l2, loc='upper center', bbox_to_anchor=(0.5, -0.12), ncol=2, fontsize=9, frameon=False)

    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format='png', bbox_inches='tight', transparent=False)
    plt.close(fig)
    out.seek(0)
    return out


def _gerar_imagem_tabela_25_anos(variaveis):
    """Gera tabela visual resumida da projeção para 25 anos."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        matplotlib.set_loglevel('warning')
        logging.getLogger('matplotlib').setLevel(logging.WARNING)
        import matplotlib.pyplot as plt
    except Exception:
        return None

    anual, acumulada = _serie_economia_25_anos_variaveis(variaveis)
    if max(acumulada) <= 0:
        return None

    anos_idx = [0, 1, 2, 4, 9, 14, 19, 24]  # anos 1,2,3,5,10,15,20,25
    linhas = []
    for i in anos_idx:
        ano = i + 1
        linhas.append([str(ano), _fmt_rs(anual[i]), _fmt_rs(acumulada[i])])

    colunas = ['Ano', 'Economia anual', 'Economia acumulada']

    fig, ax = plt.subplots(figsize=(8.6, 3.6), dpi=170)
    fig.patch.set_facecolor('#FFFFFF')
    ax.axis('off')

    tabela = ax.table(cellText=linhas, colLabels=colunas, loc='center', cellLoc='center')
    tabela.auto_set_font_size(False)
    tabela.set_fontsize(9)
    tabela.scale(1, 1.23)

    for c_idx in range(len(colunas)):
        cell = tabela[(0, c_idx)]
        cell.set_facecolor('#0F2F53')
        cell.set_edgecolor('#28507A')
        cell.set_text_props(weight='bold', color='#E9F6FF')

    for r_idx in range(1, len(linhas) + 1):
        bg = '#FFFFFF' if r_idx % 2 == 1 else '#EEF6FF'
        for c_idx in range(len(colunas)):
            cell = tabela[(r_idx, c_idx)]
            cell.set_facecolor(bg)
            cell.set_edgecolor('#E3EAF2')

    ax.set_title('Resumo de economia projetada - 25 anos', loc='left', fontsize=12, fontweight='bold', color='#0E2C4D', pad=10)

    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format='png', bbox_inches='tight', transparent=False)
    plt.close(fig)
    out.seek(0)
    return out


def _gerar_imagem_grafico_payback(variaveis):
    """Gera gráfico visual da curva de payback (investimento x economia acumulada)."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        matplotlib.set_loglevel('warning')
        logging.getLogger('matplotlib').setLevel(logging.WARNING)
        import matplotlib.pyplot as plt
        import matplotlib.patheffects as pe
    except Exception:
        return None

    investimento = _to_float_text(
        variaveis.get('valor_total')
        or variaveis.get('orcamento_valor_nota')
        or variaveis.get('VALOR_INVESTIMENTO')
        or 0,
        0.0,
    )
    economia_mensal = _to_float_text(
        variaveis.get('economia_mensal')
        or variaveis.get('ECONOMIA_MENSAL')
        or 0,
        0.0,
    )

    if economia_mensal <= 0:
        payback_anos = _extract_first_float(
            variaveis.get('payback_ano_lei_14300')
            or variaveis.get('payback_roi_lei_14300')
            or variaveis.get('PAYBACK')
            or variaveis.get('payback')
            or variaveis.get('payback_anos')
            or 0,
            0.0,
        )
        if investimento > 0 and payback_anos > 0:
            economia_mensal = investimento / (payback_anos * 12.0)

    if economia_mensal <= 0:
        fatura_sem = _to_float_text(variaveis.get('fatura_media_mensal_sem_sistema'), 0)
        fatura_com = _to_float_text(variaveis.get('fatura_minima'), 0)
        if fatura_sem > 0:
            economia_mensal = max(fatura_sem - fatura_com, fatura_sem * 0.85)

    if investimento <= 0 or economia_mensal <= 0:
        return None

    payback_meses = max(1, int(round(investimento / economia_mensal)))
    horizonte = max(36, payback_meses + 24)
    meses = list(range(0, horizonte + 1))
    acumulada = [economia_mensal * m for m in meses]

    fig, ax = plt.subplots(figsize=(11.2, 4.8), dpi=190)
    fig.patch.set_facecolor('#FFFFFF')
    ax.set_facecolor('#F7FBFF')

    line_pay, = ax.plot(meses, acumulada, color='#16B36A', linewidth=2.5, label='Economia acumulada (R$)')
    ax.fill_between(meses, acumulada, color='#BAF2D5', alpha=0.38)
    line_pay.set_path_effects([pe.Stroke(linewidth=3.8, foreground='#8EE9BE', alpha=0.85), pe.Normal()])

    ax.axhline(investimento, color='#E64A63', linewidth=1.8, linestyle='--', label='Investimento inicial (R$)')
    ax.axvline(payback_meses, color='#0A84D0', linewidth=1.6, linestyle=':', label=f'Payback: {payback_meses} meses')

    ax.scatter([payback_meses], [investimento], color='#0A84D0', edgecolors='white', linewidths=0.8, s=45, zorder=4)
    ax.text(
        payback_meses,
        investimento,
        f'  Retorno em {payback_meses} meses',
        fontsize=9,
        color='#0A84D0',
        va='bottom',
        ha='left'
    )

    ax.set_xlabel('Meses', fontsize=10, color='#24384D', fontweight='bold')
    ax.set_ylabel('Valor acumulado (R$)', fontsize=10, color='#24384D', fontweight='bold')
    ax.tick_params(axis='x', labelsize=10)
    ax.tick_params(axis='y', labelsize=10)
    ax.set_xlim(0, horizonte)
    ax.set_ylim(0, max(acumulada[-1], investimento) * 1.12)
    ax.grid(axis='y', color='#DFE7EF', linewidth=0.8, alpha=0.9)
    ax.set_axisbelow(True)

    for side in ['top', 'right']:
        ax.spines[side].set_visible(False)
    ax.spines['left'].set_color('#C9D6E2')
    ax.spines['bottom'].set_color('#C9D6E2')

    ax.set_title('Curva de payback do investimento', loc='left', fontsize=14, fontweight='bold', color='#0E2C4D', pad=12)
    ax.text(
        0.01,
        0.98,
        f'Investimento: {_fmt_rs(investimento)}  |  Economia mensal: {_fmt_rs(economia_mensal)}',
        transform=ax.transAxes,
        ha='left',
        va='top',
        fontsize=9,
        color='#1F3A56',
        bbox=dict(boxstyle='round,pad=0.25', fc='#EEF5FF', ec='#C9DDF7', lw=0.7)
    )

    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.12), ncol=3, fontsize=9, frameon=False)

    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format='png', bbox_inches='tight', transparent=False)
    plt.close(fig)
    out.seek(0)
    return out


def _substituir_placeholders_graficos_doc(doc, variaveis):
    """Substitui placeholders de gráfico por imagens reais quando o placeholder está sozinho."""
    img_12m = _gerar_imagem_grafico_consumo_geracao_12m(variaveis)
    img_tab_12m = _gerar_imagem_tabela_12m(variaveis)
    img_25a = _gerar_imagem_grafico_economia_25_anos(variaveis)
    img_tab_25a = _gerar_imagem_tabela_25_anos(variaveis)
    img_payback = _gerar_imagem_grafico_payback(variaveis)

    # Se a projeção de 25 anos não estiver disponível, usa visual de 12 meses
    # para evitar fallback textual com tabela/CSV no documento final.
    if img_25a is None and img_12m is not None:
        img_25a = io.BytesIO(img_12m.getvalue())
    if img_tab_25a is None and img_tab_12m is not None:
        img_tab_25a = io.BytesIO(img_tab_12m.getvalue())

    if img_12m is None and img_tab_12m is None and img_25a is None and img_tab_25a is None and img_payback is None:
        return

    ph_12m = _placeholder_variantes('grafico_consumo_geracao_12_meses')
    ph_tab_12m = _placeholder_variantes('tabela_12_meses') | _placeholder_variantes('tabela_12_mses')
    ph_25a = _placeholder_variantes('grafico_consumo_geracao_25_anos')
    ph_tab_25a = _placeholder_variantes('tabela_25_anos')
    ph_payback = _placeholder_variantes('grafico_pay_back') | _placeholder_variantes('grafico_payback')

    def _aplicar_paragrafo(paragrafo):
        txt = (paragrafo.text or '').strip()
        if not txt:
            return

        marcadores = [
            # (placeholders, imagem, largura, fallback_texto)
            (ph_12m, img_12m, 6.2, None),
            (
                ph_tab_12m,
                img_tab_12m,
                6.5,
                str(
                    variaveis.get('tabela_12_meses')
                    or variaveis.get('TABELA_12_MESES')
                    or variaveis.get('tabela_12_mses')
                    or variaveis.get('TABELA_12_MSES')
                    or ''
                ),
            ),
            (ph_25a, img_25a, 6.5, None),
            (
                ph_tab_25a,
                img_tab_25a,
                6.1,
                str(
                    variaveis.get('tabela_25_anos')
                    or variaveis.get('TABELA_25_ANOS')
                    or ''
                ),
            ),
            (ph_payback, img_payback, 6.5, None),
        ]

        # Troca também quando há mais de um placeholder no mesmo parágrafo.
        # Exemplo: "[grafico_25]\n[tabela_25]".
        texto_limpo = txt
        encontrou_marcador = False
        for ph_set, _, _, _ in marcadores:
            for ph in ph_set:
                if ph in texto_limpo:
                    encontrou_marcador = True
                    texto_limpo = texto_limpo.replace(ph, '')

        if not encontrou_marcador:
            return

        # Se houver texto além dos placeholders, não remove o conteúdo do usuário.
        if texto_limpo.replace('\n', '').replace('\r', '').strip():
            return

        for run in paragrafo.runs:
            run.text = ''
        run = paragrafo.runs[0] if paragrafo.runs else paragrafo.add_run()

        inseriu = False
        for ph_set, img, width, fallback_texto in marcadores:
            if not any(ph in txt for ph in ph_set):
                continue

            if img is not None:
                if inseriu:
                    run.add_break()
                run.add_picture(io.BytesIO(img.getvalue()), width=Inches(width))
                inseriu = True
                continue

            # Se a imagem não puder ser gerada, não deixa o bloco sumir.
            if fallback_texto:
                if inseriu:
                    run.add_break()
                run.add_text(fallback_texto)
                inseriu = True

    for paragrafo in doc.paragraphs:
        _aplicar_paragrafo(paragrafo)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _aplicar_paragrafo(paragrafo)

    for secao in doc.sections:
        for paragrafo in secao.header.paragraphs:
            _aplicar_paragrafo(paragrafo)
        for tabela in secao.header.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        _aplicar_paragrafo(paragrafo)

        for paragrafo in secao.footer.paragraphs:
            _aplicar_paragrafo(paragrafo)
        for tabela in secao.footer.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        _aplicar_paragrafo(paragrafo)


def _serie_consumo_12_meses(projeto):
    """Retorna lista com 12 valores de consumo mensal."""
    base = _to_float(getattr(projeto, 'consumo_kwh_mes', 0), 0)
    historico = getattr(projeto, 'historico_consumo', None)
    meses = ['Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']

    if isinstance(historico, dict):
        valores = []
        for m in meses:
            # suporta chaves em formatos diferentes
            v = historico.get(m)
            if v is None:
                v = historico.get(m.lower())
            if v is None:
                v = historico.get(m.upper())
            valores.append(_to_float(v, base))
        if any(valores):
            return valores

    return [base for _ in range(12)]


def _serie_geracao_12_meses(projeto):
    """Retorna lista com 12 valores de geração mensal com sazonalidade simples."""
    base = _to_float(getattr(projeto, 'geracao_estimada_mes', 0), 0)
    fatores = [0.93, 0.92, 0.95, 0.98, 1.02, 1.06, 1.08, 1.07, 1.03, 1.00, 0.97, 0.94]
    return [base * f for f in fatores]


def _montar_tabela_12_meses_texto(projeto):
    meses = ['Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']
    consumo = _serie_consumo_12_meses(projeto)
    geracao = _serie_geracao_12_meses(projeto)

    linhas = [
        'SOLAR GRID :: 12M PERFORMANCE',
        'Mes | Consumo(kWh) | Geracao(kWh) | Saldo(kWh) | Nivel',
        '----+--------------+--------------+------------+----------------',
    ]
    pico = max(max(consumo), max(geracao), 1)
    for i, mes in enumerate(meses):
        saldo = geracao[i] - consumo[i]
        barra = _bar(geracao[i], pico)
        linhas.append(f"{mes:>3} | {consumo[i]:>12.0f} | {geracao[i]:>12.0f} | {saldo:>10.0f} | {barra}")

    total_c = sum(consumo)
    total_g = sum(geracao)
    total_s = total_g - total_c
    linhas.append('----+--------------+--------------+------------+----------------')
    linhas.append(f"TOT | {total_c:>12.0f} | {total_g:>12.0f} | {total_s:>10.0f} |")
    return '\n'.join(linhas)


def _montar_tabela_25_anos_texto(projeto):
    economia_mensal = _to_float(getattr(projeto, 'economia_mensal', 0), 0)
    reajuste = _to_float(
        getattr(projeto, 'acrescimo_anual_percentual', None)
        or getattr(projeto, 'reajuste_anual_energia', None)
        or 10.0,
        10.0,
    )
    # Interpretação: campo vem em %, usamos como crescimento de tarifa/economia para projeção.
    taxa = reajuste / 100.0

    linhas = [
        'ECONOMIC TIMELINE :: 25Y FORECAST',
        'Ano | Economia Anual (R$) | Acumulado (R$) | Evolucao',
        '----+----------------------+----------------+----------------',
    ]
    acumulado = 0.0
    economia_ano_base = economia_mensal * 12
    maior_anual = economia_ano_base * ((1 + taxa) ** 24) if economia_ano_base else 1
    for ano in range(1, 26):
        economia_ano = economia_ano_base * ((1 + taxa) ** (ano - 1))
        acumulado += economia_ano
        barra = _bar(economia_ano, maior_anual)
        linhas.append(f"{ano:>3} | {economia_ano:>20.2f} | {acumulado:>14.2f} | {barra}")
    return '\n'.join(linhas)


def _montar_grafico_12_meses_csv(projeto):
    meses = ['Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']
    consumo = _serie_consumo_12_meses(projeto)
    geracao = _serie_geracao_12_meses(projeto)

    linhas = [
        'CHART DATA :: CONSUMO x GERACAO (12M)',
        'Mes | Consumo | Geracao | Delta | C-Bar             | G-Bar',
        '----+---------+---------+-------+-------------------+-------------------',
    ]
    pico = max(max(consumo), max(geracao), 1)
    for i, mes in enumerate(meses):
        delta = geracao[i] - consumo[i]
        c_bar = _bar(consumo[i], pico, largura=19)
        g_bar = _bar(geracao[i], pico, largura=19)
        linhas.append(f"{mes:>3} | {consumo[i]:>7.0f} | {geracao[i]:>7.0f} | {delta:>5.0f} | {c_bar} | {g_bar}")
    return '\n'.join(linhas)


def _montar_grafico_25_anos_csv(projeto):
    economia_mensal = _to_float(getattr(projeto, 'economia_mensal', 0), 0)
    reajuste = _to_float(
        getattr(projeto, 'acrescimo_anual_percentual', None)
        or getattr(projeto, 'reajuste_anual_energia', None)
        or 10.0,
        10.0,
    )
    taxa = reajuste / 100.0

    if economia_mensal <= 0:
        return 'Dados insuficientes para projeção de economia em 25 anos.'

    linhas = [
        'CHART DATA :: ECONOMIA 25 ANOS',
        'Ano | Economia_Anual | Economia_Acumulada | Curva',
        '----+----------------+--------------------+----------------',
    ]
    acumulado = 0.0
    economia_ano_base = economia_mensal * 12
    maior_anual = economia_ano_base * ((1 + taxa) ** 24) if economia_ano_base else 1
    for ano in range(1, 26):
        economia_ano = economia_ano_base * ((1 + taxa) ** (ano - 1))
        acumulado += economia_ano
        linha = _bar(economia_ano, maior_anual)
        linhas.append(f"{ano:>3} | {economia_ano:>14.2f} | {acumulado:>18.2f} | {linha}")
    return '\n'.join(linhas)


def _montar_grafico_payback_csv(projeto):
    investimento = _to_float(
        getattr(projeto, 'valor_venda', None)
        or getattr(projeto, 'valor_orcamento', None)
        or getattr(projeto, 'custo_total', None)
        or 0,
        0,
    )
    economia_mensal = _to_float(getattr(projeto, 'economia_mensal', 0), 0)
    payback_anos = _to_float(
        getattr(projeto, 'payback_anos', None)
        or getattr(projeto, 'payback', None)
        or 0,
        0,
    )

    if economia_mensal <= 0 and investimento > 0 and payback_anos > 0:
        economia_mensal = investimento / (payback_anos * 12.0)

    if economia_mensal <= 0:
        return 'Dados insuficientes para curva de payback.'
    meses = max(24, int(payback_anos * 12) + 12)

    payback_mes_estimado = int(round(investimento / economia_mensal)) if economia_mensal > 0 else 0

    linhas = [
        'CHART DATA :: CURVA DE PAYBACK',
        f'Investimento (R$): {investimento:.2f} | Economia mensal (R$): {economia_mensal:.2f} | Payback estimado: {payback_mes_estimado} meses',
        'Mes | Investimento | Economia_Acumulada | Progresso',
        '----+-------------+--------------------+----------------',
    ]
    for m in range(0, meses + 1, 2):
        economia_acumulada = economia_mensal * m
        progresso = _bar(economia_acumulada, max(investimento, economia_acumulada, 1))
        linhas.append(f"{m:>3} | {investimento:>11.2f} | {economia_acumulada:>18.2f} | {progresso}")
    return '\n'.join(linhas)


def _substituir_texto_paragrafo(paragrafo, variaveis):
    """Substitui placeholders em múltiplos formatos no parágrafo."""
    if not paragrafo.text:
        return

    texto_original = paragrafo.text
    texto_novo = texto_original

    for variavel, valor in variaveis.items():
        chave = str(variavel)
        valor_str = str(valor)

        placeholders = [
            f'[{chave}]',
            f'[{chave.upper()}]',
            f'[{chave.lower()}]',
            f'{{{{{chave}}}}}',
            f'{{{{{chave.upper()}}}}}',
            f'{{{{{chave.lower()}}}}}',
            f'{{{chave}}}',
            f'{{{chave.upper()}}}',
            f'{{{chave.lower()}}}',
            f'<<{chave}>>',
            f'<<{chave.upper()}>>',
            f'<<{chave.lower()}>>',
            f'${{{chave}}}',
            f'${{{chave.upper()}}}',
            f'${{{chave.lower()}}}',
            f'%{chave}%',
            f'%{chave.upper()}%',
            f'%{chave.lower()}%',
        ]

        for ph in placeholders:
            if ph in texto_novo:
                texto_novo = texto_novo.replace(ph, valor_str)

    if texto_novo != texto_original:
        for run in paragrafo.runs:
            run.text = ''
        if paragrafo.runs:
            paragrafo.runs[0].text = texto_novo
        else:
            paragrafo.add_run(texto_novo)


def _substituir_placeholders_xml_docx(caminho_docx, variaveis):
    """
    Fallback robusto no XML interno do DOCX.
    Cobre text boxes/shapes e estruturas que python-docx não percorre bem.
    """
    base_dir = os.path.dirname(os.path.abspath(caminho_docx)) or None
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=base_dir) as tmp:
        tmp_saida = tmp.name

    try:
        with zipfile.ZipFile(caminho_docx, 'r') as zin, zipfile.ZipFile(tmp_saida, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                    xml = data.decode('utf-8')
                    xml_novo = xml

                    for variavel, valor in variaveis.items():
                        chave = str(variavel)
                        valor_str = str(valor)

                        # Para blocos multiline (tabelas/graficos), deixamos python-docx
                        # aplicar no nível de parágrafo para preservar melhor as quebras.
                        if '\n' in valor_str or chave.lower().startswith('grafico_'):
                            continue

                        valor_xml = escape(valor_str)

                        placeholders = [
                            f'[{chave}]', f'[{chave.upper()}]', f'[{chave.lower()}]',
                            f'{{{{{chave}}}}}', f'{{{{{chave.upper()}}}}}', f'{{{{{chave.lower()}}}}}',
                            f'{{{chave}}}', f'{{{chave.upper()}}}', f'{{{chave.lower()}}}',
                            f'<<{chave}>>', f'<<{chave.upper()}>>', f'<<{chave.lower()}>>',
                            f'${{{chave}}}', f'${{{chave.upper()}}}', f'${{{chave.lower()}}}',
                            f'%{chave}%', f'%{chave.upper()}%', f'%{chave.lower()}%',
                        ]

                        for ph in placeholders:
                            if ph in xml_novo:
                                xml_novo = xml_novo.replace(ph, valor_xml)

                    if xml_novo != xml:
                        data = xml_novo.encode('utf-8')

                zout.writestr(item, data)

        # Mesmo diretório evita erro de cross-device em ambientes como Render
        os.replace(tmp_saida, caminho_docx)
    finally:
        if 'tmp_saida' in locals() and os.path.exists(tmp_saida):
            os.remove(tmp_saida)


def _expandir_aliases_variaveis(variaveis):
    """Gera aliases de chave em maiúsculas e minúsculas para máxima compatibilidade."""
    expandidas = {}
    for chave, valor in variaveis.items():
        c = str(chave)
        expandidas[c] = valor
        expandidas[c.upper()] = valor
        expandidas[c.lower()] = valor

    # Aliases legados usados em modelos antigos
    if 'NOME_CLIENTE' in expandidas:
        expandidas['cliente_nome'] = expandidas['NOME_CLIENTE']
    if 'CPF_CNPJ_CLIENTE' in expandidas:
        expandidas['cliente_cpf_cnpj'] = expandidas['CPF_CNPJ_CLIENTE']
    if 'NUMERO_PROJETO' in expandidas:
        expandidas['id_projeto'] = expandidas['NUMERO_PROJETO']
        expandidas['numero_projeto'] = expandidas['NUMERO_PROJETO']
    if 'DATA_PROPOSTA' in expandidas:
        expandidas['data_proposta'] = expandidas['DATA_PROPOSTA']

    return expandidas


def _coletar_campos_simples(obj):
    """Coleta atributos escalares simples de um objeto SQLAlchemy/model."""
    if not obj:
        return {}

    dados = {}
    for chave, valor in vars(obj).items():
        if chave.startswith('_'):
            continue
        if isinstance(valor, (str, int, float, bool)) or valor is None:
            dados[chave] = '' if valor is None else valor
    return dados


def substituir_variaveis_word(template_path, variaveis):
    """
    Substitui variáveis em um template Word
    
    Args:
        template_path: Caminho do arquivo .docx template
        variaveis: Dict com {variavel: valor} para substituição
        
    Returns:
        Document object com as substituições feitas
    """
    variaveis = _expandir_aliases_variaveis(variaveis)
    variaveis = _normalizar_variaveis_financeiras(variaveis)
    variaveis = _expandir_aliases_variaveis(variaveis)

    doc = Document(template_path)

    # Troca placeholders de gráficos por imagens reais quando possível.
    # Deve ocorrer antes da substituição textual para não perder o marcador.
    _substituir_placeholders_graficos_doc(doc, variaveis)
    
    # Substituir em parágrafos
    for paragrafo in doc.paragraphs:
        _substituir_texto_paragrafo(paragrafo, variaveis)
    
    # Substituir em tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_texto_paragrafo(paragrafo, variaveis)
    
    # Substituir em cabeçalhos e rodapés (inclui primeira página e pares/ímpares)
    for secao in doc.sections:
        blocos_hf = [
            secao.header,
            secao.first_page_header,
            secao.even_page_header,
            secao.footer,
            secao.first_page_footer,
            secao.even_page_footer,
        ]

        for bloco in blocos_hf:
            for paragrafo in bloco.paragraphs:
                _substituir_texto_paragrafo(paragrafo, variaveis)
            for tabela in bloco.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            _substituir_texto_paragrafo(paragrafo, variaveis)
    
    return doc


def gerar_variaveis_projeto(projeto, cliente=None, config=None, balanco=None):
    """
    Gera dicionário de variáveis a partir de um ProjetoSolar
    
    Args:
        projeto: Instância de ProjetoSolar
        cliente: Instância de Cliente (opcional)
        config: Configuração da empresa
        balanco: Dict com balanço energético
        
    Returns:
        Dict com todas as variáveis disponíveis
    """
    from app.energia_solar.catalogo_model import KitSolar, PlacaSolar, InversorSolar
    
    # Carregar kit/placa/inversor
    kit = getattr(projeto, 'kit', None)
    if not kit and getattr(projeto, 'kit_id', None):
        kit = KitSolar.query.get(projeto.kit_id)

    placa = PlacaSolar.query.get(projeto.placa_id) if projeto.placa_id else None
    inversor = InversorSolar.query.get(projeto.inversor_id) if projeto.inversor_id else None

    # Fallback: quando projeto usa modo kit, os componentes podem não estar em placa_id/inversor_id.
    if not placa and kit and getattr(kit, 'placa', None):
        placa = kit.placa
    if not inversor and kit and getattr(kit, 'inversor', None):
        inversor = kit.inversor

    placa_modelo = (
        (placa.modelo if placa else None)
        or getattr(projeto, 'placa_modelo', None)
        or getattr(projeto, 'modulo_modelo', None)
        or getattr(projeto, 'modelo_modulo', None)
        or ''
    )
    placa_fabricante = (
        (placa.fabricante if placa else None)
        or getattr(projeto, 'placa_fabricante', None)
        or getattr(projeto, 'fabricante_modulo', None)
        or ''
    )
    placa_potencia_val = (
        (placa.potencia if placa and getattr(placa, 'potencia', None) else None)
        or getattr(projeto, 'placa_potencia', None)
        or getattr(projeto, 'potencia_modulo', None)
        or getattr(projeto, 'modulo_potencia', None)
    )
    placa_potencia_num = _extract_first_float(placa_potencia_val, 0)
    placa_potencia = f"{placa_potencia_num:.0f}W" if placa_potencia_num > 0 else ''

    inversor_modelo = (
        (inversor.modelo if inversor else None)
        or getattr(projeto, 'inversor_modelo', None)
        or ''
    )
    inversor_fabricante = (
        (inversor.fabricante if inversor else None)
        or getattr(projeto, 'inversor_fabricante', None)
        or ''
    )
    inversor_potencia_val = (
        (inversor.potencia_nominal if inversor and hasattr(inversor, 'potencia_nominal') else None)
        or (inversor.potencia_maxima if inversor and hasattr(inversor, 'potencia_maxima') else None)
        or getattr(projeto, 'inversor_potencia', None)
        or getattr(projeto, 'potencia_inversor', None)
    )
    inversor_potencia_num = _extract_first_float(inversor_potencia_val, 0)
    inversor_potencia = f"{inversor_potencia_num:.1f}kW" if inversor_potencia_num > 0 else ''
    
    variaveis = {
        # Projeto
        'id_projeto': projeto.id or '',
        'projeto_titulo': projeto.nome_cliente or '',
        'nome_cliente': projeto.nome_cliente or '',
        'data_criacao': projeto.data_criacao.strftime('%d/%m/%Y') if projeto.data_criacao else '',
        'status': projeto.status or 'rascunho',
        
        # Endereço
        'endereco': projeto.endereco or '',
        'cidade': projeto.cidade or '',
        'estado': projeto.estado or '',
        'cep': projeto.cep or '',
        'latitude': projeto.latitude or '',
        'longitude': projeto.longitude or '',
        
        # Irradiação
        'irradiacao_solar_media': f"{projeto.irradiacao_solar:.2f}" if projeto.irradiacao_solar else '',
        
        # Sistema
        'qtd_placas': projeto.qtd_placas or 0,
        'potencia_kwp': f"{projeto.potencia_kwp:.2f}" if projeto.potencia_kwp else '',
        'geracao_estimada_mes': f"{projeto.geracao_estimada_mes:.0f}" if projeto.geracao_estimada_mes else '',
        'area_ocupada': f"{(projeto.qtd_placas or 0) * 2.5:.1f}",
        'peso_total': f"{(projeto.qtd_placas or 0) * 25:.0f}",
        
        # Consumo
        'consumo_kwh_mes': f"{projeto.consumo_kwh_mes:.0f}" if projeto.consumo_kwh_mes else '',
        'simultaneidade': f"{projeto.simultaneidade:.0f}" if projeto.simultaneidade else '35',
        'tarifa_kwh': f"{projeto.tarifa_kwh:.2f}" if projeto.tarifa_kwh else '',
        'tipo_instalacao': (projeto.tipo_instalacao or 'monofasica').capitalize(),
        
        # Placas
        'placa_fabricante': placa_fabricante,
        'placa_modelo': placa_modelo,
        'placa_potencia': placa_potencia,
        'placa_tensao': f"{placa.tensao_maxima_potencia}V" if placa and hasattr(placa, 'tensao_maxima_potencia') and placa.tensao_maxima_potencia else '',
        'placa_corrente': f"{placa.corrente_maxima_potencia}A" if placa and hasattr(placa, 'corrente_maxima_potencia') and placa.corrente_maxima_potencia else '',
        'placa_eficiencia': f"{placa.eficiencia}%" if placa and hasattr(placa, 'eficiencia') and placa.eficiencia else '',
        'placa_garantia': '25 anos',
        'placa_tecnologia': placa.tecnologia if placa and hasattr(placa, 'tecnologia') else '',
        
        # Inversor
        'inversor_fabricante': inversor_fabricante,
        'inversor_modelo': inversor_modelo,
        'inversor_potencia': inversor_potencia,
        'inversor_tensao_entrada': f"{inversor.tensao_mppt_min}-{inversor.tensao_mppt_max}V" if inversor and hasattr(inversor, 'tensao_mppt_min') and inversor.tensao_mppt_min else '',
        'inversor_tensao_saida': inversor.tensao_saida if inversor and hasattr(inversor, 'tensao_saida') else '220V',
        'inversor_eficiencia': (
            f"{inversor.eficiencia_maxima}%"
            if inversor and hasattr(inversor, 'eficiencia_maxima') and inversor.eficiencia_maxima
            else ''
        ),
        'inversor_garantia': '10 anos',
        'inversor_monitoramento': 'WiFi/App',
        
        # Valores
        'valor_total': f"{projeto.valor_venda:.2f}" if projeto.valor_venda else '',
        'valor_vista': f"{projeto.valor_venda * 0.95:.2f}" if projeto.valor_venda else '',
        
        # Empresa
        'empresa_nome': config.nome_fantasia if config else 'JSP Elétrica & Solar',
        'empresa_razao': config.razao_social if config else 'JSP Elétrica Industrial Ltda',
        'empresa_cnpj': config.cnpj if config else '41.280.764/0001-65',
        'empresa_endereco': config.logradouro if config else 'Rua Indalécio Costa, 890',
        'empresa_cidade': config.cidade if config else 'Tietê',
        'empresa_estado': config.uf if config else 'SP',
        'empresa_cep': config.cep if config else '18530-170',
        'empresa_telefone': config.telefone if config else '(15) 99670-2036',
        'empresa_email': config.email if config else 'atendimento@eletricasaroba.com.br',
        'empresa_site': 'www.eletricasaroba.com.br',
    }

    # Variáveis extras encontradas em modelos comerciais personalizados
    valor_conta_luz = getattr(projeto, 'valor_conta_luz', None)
    consumo_kwh = getattr(projeto, 'consumo_kwh_mes', None)
    tarifa_kwh = getattr(projeto, 'tarifa_kwh', None)
    fatura_media = 0

    for fonte in [
        valor_conta_luz,
        getattr(projeto, 'fatura_media_mensal_sem_sistema', None),
        getattr(projeto, 'fatura_sem_sistema', None),
        getattr(projeto, 'conta_luz_atual', None),
    ]:
        fatura_media = _extract_first_float(fonte, 0)
        if fatura_media > 0:
            break

    if fatura_media <= 0 and consumo_kwh and tarifa_kwh:
        try:
            fatura_media = float(consumo_kwh) * float(tarifa_kwh)
        except Exception:
            fatura_media = 0

    # Economia mensal robusta para evitar zeros indevidos em gráficos/tabelas.
    economia_mensal_calc = _to_float(
        getattr(projeto, 'economia_mensal', None)
        or getattr(projeto, 'economia', None)
        or 0,
        0,
    )

    if economia_mensal_calc <= 0:
        if fatura_media > 0:
            economia_mensal_calc = fatura_media * 0.9
        elif consumo_kwh and tarifa_kwh:
            try:
                economia_mensal_calc = float(consumo_kwh) * float(tarifa_kwh) * 0.9
            except Exception:
                economia_mensal_calc = 0

    # Se a economia existe mas a fatura não foi informada, estima fatura sem sistema.
    if fatura_media <= 0 and economia_mensal_calc > 0:
        fatura_media = economia_mensal_calc / 0.9

    tipo_instalacao = getattr(projeto, 'tipo_instalacao', None) or 'monofasica'
    consumo_minimo_kwh = _consumo_minimo_kwh_por_tipo(tipo_instalacao)

    adicionais = 0.0
    for chave in [
        'iluminacao_publica',
        'demais_custos',
        'outros_custos',
        'adicionais_fatura',
        'cip',
        'taxa_disponibilidade',
        'taxa_iluminacao_publica',
    ]:
        adicionais += _extract_first_float(getattr(projeto, chave, None), 0)

    tarifa_num = _extract_first_float(tarifa_kwh, 0)
    if tarifa_num <= 0 and fatura_media > 0 and consumo_kwh:
        tarifa_num = _extract_first_float(fatura_media / max(float(consumo_kwh), 1.0), 0)
    if tarifa_num <= 0:
        tarifa_num = 0.85

    fatura_minima_tecnica = (consumo_minimo_kwh * tarifa_num) + adicionais

    fatura_com_sistema = _extract_first_float(getattr(projeto, 'fatura_com_sistema', None), 0)
    if fatura_com_sistema <= 0 and fatura_media > 0:
        fatura_com_sistema = max(fatura_media - economia_mensal_calc, fatura_minima_tecnica)

    if fatura_com_sistema > 0:
        fatura_com_sistema = max(fatura_com_sistema, fatura_minima_tecnica)

    if fatura_media > 0 and fatura_media < fatura_minima_tecnica:
        fatura_media = fatura_minima_tecnica

    if fatura_com_sistema > fatura_media > 0:
        fatura_com_sistema = fatura_media

    economia_anual_calc = economia_mensal_calc * 12
    economia_25_calc = economia_anual_calc * 25

    valor_nota = (
        getattr(projeto, 'valor_nota_fiscal', None)
        or getattr(projeto, 'valor_venda', None)
        or getattr(projeto, 'valor_orcamento', None)
        or getattr(projeto, 'custo_total', None)
        or 0
    )

    kit_desc = ''
    kit_obj = getattr(projeto, 'kit', None)
    if kit_obj:
        kit_desc = (
            getattr(kit_obj, 'descricao', None)
            or getattr(kit_obj, 'nome', None)
            or getattr(kit_obj, 'modelo', None)
            or ''
        )

    kit_outras_inf = (
        (getattr(kit_obj, 'outras_informacoes', None) if kit_obj else None)
        or ''
    )

    outras_descricoes = (
        kit_outras_inf
        or getattr(projeto, 'observacoes', None)
        or getattr(projeto, 'descricao', None)
        or ''
    )

    payback_anos = (
        getattr(projeto, 'payback_anos', None)
        or getattr(projeto, 'payback', None)
        or 0
    )

    variaveis.update({
        'fatura_media_mensal_sem_sistema': f"{float(fatura_media or 0):.2f}",
        'fatura_minima': f"{float(fatura_com_sistema or 0):.2f}",
        'fatura_minima_tecnica': f"{float(fatura_minima_tecnica or 0):.2f}",
        'fatura_sem_sistema': f"{float(fatura_media or 0):.2f}",
        'fatura_com_sistema': f"{float(fatura_com_sistema or 0):.2f}",
        'fatura_com_sistema_rs': _fmt_rs(fatura_com_sistema),
        'fatura_sem_sistema_rs': _fmt_rs(fatura_media),
        'conta_luz_atual': f"{float(fatura_media or 0):.2f}",
        'conta_luz_futura': f"{float(fatura_com_sistema or 0):.2f}",
        'conta_luz_atual_rs': _fmt_rs(fatura_media),
        'conta_luz_futura_rs': _fmt_rs(fatura_com_sistema),
        'consumo_minimo_kwh': f"{float(consumo_minimo_kwh or 0):.0f}",
        'economia_mensal': f"{float(economia_mensal_calc or 0):.2f}",
        'economia_anual': f"{float(economia_anual_calc or 0):.2f}",
        'economia_25_anos': f"{float(economia_25_calc or 0):.2f}",
        'economia_total_25_anos': f"{float(economia_25_calc or 0):.2f}",
        # Placeholders visuais devem ser preenchidos por imagem, não por texto ASCII.
        'grafico_consumo_geracao_12_meses': '',
        'grafico_consumo_geracao_25_anos': '',
        'grafico_pay_back': '',
        'kit_descricao': kit_desc,
        'kit_outras_inf': outras_descricoes,
        'kit_outras_informacoes': outras_descricoes,
        'outras_informacoes_kit': outras_descricoes,
        'outras_informacoes': outras_descricoes,
        # Aliases para modelos que usam nomenclatura diferente
        'outras_descricoes': outras_descricoes,
        'outras_descricoes_kit': outras_descricoes,
        'outras_descricao': outras_descricoes,
        'descricao_complementar': outras_descricoes,
        'orcamento_valor_nota': f"{float(valor_nota or 0):.2f}",
        'payback_ano_lei_14300': f"{float(payback_anos or 0):.1f}",
        'payback_roi_lei_14300': f"{float(payback_anos or 0):.1f}",
        'reajuste_anual': f"{float(getattr(projeto, 'acrescimo_anual_percentual', None) or getattr(projeto, 'reajuste_anual_energia', None) or 10.0):.2f}",
        'acrescimo_anual_percentual': f"{float(getattr(projeto, 'acrescimo_anual_percentual', None) or getattr(projeto, 'reajuste_anual_energia', None) or 10.0):.2f}",
        'tabela_12_meses': '',
        'tabela_25_anos': '',
    })
    
    # Adicionar variáveis do balanço se fornecido
    if balanco:
        variaveis.update({
            'consumo_simultaneo': f"{balanco.get('consumo_simultaneo', 0):.0f}",
            'excedente_kwh': f"{balanco.get('excedente_rede', 0):.0f}",
            'economia_mensal': f"{balanco.get('economia_mensal', 0):.2f}",
            'economia_anual': f"{balanco.get('economia_mensal', 0) * 12:.2f}",
            'economia_25_anos': f"{balanco.get('economia_mensal', 0) * 12 * 25:.2f}",
            'consumo_minimo': balanco.get('consumo_minimo', 30),
        })

    # Aliases para templates em formato comercial (maiúsculas) e fluxo legado
    variaveis.update({
        'NOME_CLIENTE': (
            (cliente.nome_razao_social if cliente and hasattr(cliente, 'nome_razao_social') else None)
            or (cliente.razao_social if cliente and hasattr(cliente, 'razao_social') else None)
            or (cliente.nome if cliente and hasattr(cliente, 'nome') else None)
            or projeto.nome_cliente
            or ''
        ),
        'NUMERO_PROJETO': projeto.id or '',
        'CPF_CNPJ_CLIENTE': cliente.cpf_cnpj if cliente and hasattr(cliente, 'cpf_cnpj') else '',
        'CIDADE': (cliente.cidade if cliente and hasattr(cliente, 'cidade') else None) or projeto.cidade or '',
        'ESTADO': (cliente.estado if cliente and hasattr(cliente, 'estado') else None) or projeto.estado or '',
    })

    # Expor campos simples automaticamente para aceitar qualquer placeholder útil no documento
    campos_projeto = _coletar_campos_simples(projeto)
    for k, v in campos_projeto.items():
        variaveis.setdefault(k, v)
        variaveis.setdefault(f'projeto_{k}', v)

    campos_cliente = _coletar_campos_simples(cliente)
    for k, v in campos_cliente.items():
        variaveis.setdefault(f'cliente_{k}', v)

    campos_config = _coletar_campos_simples(config)
    for k, v in campos_config.items():
        variaveis.setdefault(f'empresa_{k}', v)

    # Unificar com o contexto completo usado na rota de proposta comercial.
    # Isso garante placeholders como DATA_PROPOSTA, VALIDADE_PROPOSTA,
    # VALOR_INVESTIMENTO, PAYBACK etc. também no fluxo de upload.
    try:
        from app.energia_solar.proposta_word_service import montar_contexto_proposta

        contexto_completo = montar_contexto_proposta(projeto)
        if isinstance(contexto_completo, dict):
            variaveis.update(contexto_completo)
    except Exception:
        # Mantém compatibilidade mesmo que o contexto completo falhe por algum motivo.
        pass
    
    return _expandir_aliases_variaveis(variaveis)
