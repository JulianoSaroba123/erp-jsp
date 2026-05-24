"""
Serviço de geração de propostas comerciais usando modelo Word (.docx)
Substitui placeholders por dados reais do projeto solar
"""

from pathlib import Path
from datetime import datetime, timedelta
from docx import Document
import subprocess
import os
import logging
import zipfile
import tempfile
import shutil
from xml.sax.saxutils import escape
import re

logger = logging.getLogger(__name__)

# Importar modelos separadamente da configuração para não bloquear o contexto
try:
    from app.cliente.cliente_model import Cliente
    from app.energia_solar.catalogo_model import KitSolar, PlacaSolar, InversorSolar
    MODELS_IMPORTED = True
except ImportError as e:
    logger.warning(f"Não foi possível importar modelos: {e}")
    MODELS_IMPORTED = False

try:
    from app.configuracao.configuracao_utils import carregar_configuracao
    CONFIG_IMPORTED = True
except ImportError as e:
    logger.warning(f"Não foi possível importar configuração da empresa: {e}")
    CONFIG_IMPORTED = False


def formatar_moeda(valor):
    """Formata valor numérico para padrão brasileiro R$ 1.234,56"""
    try:
        valor = float(valor or 0)
        return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "R$ 0,00"


def formatar_numero(valor, casas=2):
    """Formata número com vírgula como separador decimal"""
    try:
        return f"{float(valor or 0):.{casas}f}".replace(".", ",")
    except Exception:
        return "0"


def _to_float_flex(valor, default=0.0):
    """Converte número aceitando strings com R$, vírgula e separador de milhar."""
    if valor is None:
        return default
    if isinstance(valor, (int, float)):
        return float(valor)
    try:
        txt = str(valor).strip().replace("R$", "").replace(" ", "")
        if "," in txt and "." in txt:
            txt = txt.replace(".", "").replace(",", ".")
        else:
            txt = txt.replace(",", ".")
        return float(txt)
    except Exception:
        return default


def _extract_first_float(valor, default=0.0):
    """Extrai o primeiro número de texto como float (ex.: '3,0 anos' -> 3.0)."""
    if valor is None:
        return default
    if isinstance(valor, (int, float)):
        return float(valor)
    m = re.search(r"-?\d+[\.,]?\d*", str(valor))
    if not m:
        return default
    try:
        return float(m.group(0).replace(",", "."))
    except Exception:
        return default


def _normalizar_tipo_instalacao(valor):
    txt = (str(valor or '').strip().lower())
    if 'tri' in txt:
        return 'trifasica'
    if 'bi' in txt:
        return 'bifasica'
    return 'monofasica'


def _consumo_minimo_kwh_por_tipo(tipo_instalacao):
    tipo = _normalizar_tipo_instalacao(tipo_instalacao)
    mapa = {
        'monofasica': 30.0,
        'bifasica': 50.0,
        'trifasica': 100.0,
    }
    return mapa.get(tipo, 30.0)


def substituir_texto_em_paragrafos(doc, contexto):
    """Substitui placeholders {{CHAVE}} nos parágrafos do documento"""
    for p in doc.paragraphs:
        substituir_em_paragrafo_preservando_basico(p, contexto)


def substituir_texto_em_tabelas(doc, contexto):
    """Substitui placeholders {{CHAVE}} nas tabelas do documento"""
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir_em_paragrafo_preservando_basico(p, contexto)


def substituir_texto_em_headers_footers(doc, contexto):
    """Substitui placeholders em cabeçalhos e rodapés de todas as seções."""
    for section in doc.sections:
        for p in section.header.paragraphs:
            substituir_em_paragrafo_preservando_basico(p, contexto)
        for tabela in section.header.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        substituir_em_paragrafo_preservando_basico(p, contexto)

        for p in section.footer.paragraphs:
            substituir_em_paragrafo_preservando_basico(p, contexto)
        for tabela in section.footer.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        substituir_em_paragrafo_preservando_basico(p, contexto)


def substituir_placeholders_xml_docx(docx_path, contexto):
    """
    Fallback robusto: substitui placeholders diretamente nos XMLs do .docx.
    Cobre casos como text boxes/shapes onde python-docx não alcança bem.
    """
    substituicoes = 0

    def _placeholder_variantes(chave):
        c = str(chave)
        return {
            f'[{c}]', f'[{c.upper()}]', f'[{c.lower()}]',
            f'{{{{{c}}}}}', f'{{{{{c.upper()}}}}}', f'{{{{{c.lower()}}}}}',
            f'{{{c}}}', f'{{{c.upper()}}}', f'{{{c.lower()}}}',
            f'<<{c}>>', f'<<{c.upper()}>>', f'<<{c.lower()}>>',
            f'${{{c}}}', f'${{{c.upper()}}}', f'${{{c.lower()}}}',
            f'%{c}%', f'%{c.upper()}%', f'%{c.lower()}%',
        }

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = tmp.name

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename.endswith(".xml") and item.filename.startswith("word/"):
                    texto = data.decode("utf-8")
                    texto_original = texto

                    for chave, valor in contexto.items():
                        valor_xml = escape(str(valor))
                        for placeholder in _placeholder_variantes(chave):
                            if placeholder in texto:
                                texto = texto.replace(placeholder, valor_xml)

                    if texto != texto_original:
                        data = texto.encode("utf-8")
                        substituicoes += 1

                zout.writestr(item, data)

        shutil.move(tmp_path, docx_path)
        return substituicoes
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def substituir_em_paragrafo_preservando_basico(paragrafo, contexto):
    """
    Substitui placeholders mesmo quando o Word divide o texto em múltiplos runs.
    Usa abordagem mais robusta que reconstrói o texto completo.
    """
    texto_original = paragrafo.text
    texto_novo = texto_original

    for chave, valor in contexto.items():
        texto_novo = texto_novo.replace("{{" + chave + "}}", str(valor))

    if texto_novo != texto_original:
        # Limpa todos os runs e coloca o texto novo no primeiro
        for run in paragrafo.runs:
            run.text = ""
        if paragrafo.runs:
            paragrafo.runs[0].text = texto_novo
        else:
            paragrafo.add_run(texto_novo)


def montar_contexto_proposta(projeto):
    """
    Monta dicionário com todos os dados do projeto para substituição no Word.
    Usa lógica centralizadora similar à função montar_contexto_proposta_solar.
    """
    
    # Valor total do investimento (prioridade: valor_venda > valor_orcamento > custo_total)
    valor_total = 0.0
    for fonte in [
        getattr(projeto, "valor_venda", None),
        getattr(projeto, "valor_orcamento", None),
        getattr(projeto, "valor_nota_fiscal", None),
        getattr(projeto, "custo_total", None),
    ]:
        valor_total = _to_float_flex(fonte, 0)
        if valor_total > 0:
            break

    # Economia mensal e anual
    economia_mensal = _to_float_flex(
        getattr(projeto, "economia_mensal", None)
        or getattr(projeto, "economia", None),
        0,
    )

    valor_conta_luz = _to_float_flex(
        getattr(projeto, "valor_conta_luz", None)
        or getattr(projeto, "fatura_sem_sistema", None)
        or getattr(projeto, "fatura_media_mensal_sem_sistema", None)
        or getattr(projeto, "conta_luz_atual", None),
        0,
    )

    consumo_mensal = _to_float_flex(getattr(projeto, 'consumo_kwh_mes', None), 0)
    preco_kwh = _to_float_flex(
        getattr(projeto, 'preco_kwh', None)
        or getattr(projeto, 'tarifa_energia', None)
        or getattr(projeto, 'tarifa_kwh', None),
        0,
    )

    if valor_conta_luz <= 0 and consumo_mensal > 0 and preco_kwh > 0:
        valor_conta_luz = consumo_mensal * preco_kwh

    payback_informado = _extract_first_float(
        getattr(projeto, 'payback_anos', None)
        or getattr(projeto, 'payback', None),
        0,
    )

    if economia_mensal <= 0 and payback_informado > 0 and valor_total > 0:
        economia_mensal = valor_total / (payback_informado * 12)

    if economia_mensal <= 0 and valor_conta_luz > 0:
        economia_mensal = valor_conta_luz * 0.9

    if valor_conta_luz <= 0 and economia_mensal > 0:
        valor_conta_luz = economia_mensal / 0.9

    if preco_kwh <= 0:
        preco_kwh = 0.85

    tipo_instalacao_norm = _normalizar_tipo_instalacao(getattr(projeto, 'tipo_instalacao', None) or 'monofasica')
    consumo_minimo_kwh = _consumo_minimo_kwh_por_tipo(tipo_instalacao_norm)

    adicionais_fatura = 0.0
    for chave in [
        'iluminacao_publica',
        'demais_custos',
        'outros_custos',
        'adicionais_fatura',
        'cip',
        'taxa_disponibilidade',
        'taxa_iluminacao_publica',
    ]:
        adicionais_fatura += _to_float_flex(getattr(projeto, chave, None), 0)

    fatura_minima_tecnica = (consumo_minimo_kwh * preco_kwh) + adicionais_fatura

    conta_luz_futura = _to_float_flex(getattr(projeto, 'fatura_com_sistema', None), 0)
    if conta_luz_futura <= 0 and valor_conta_luz > 0:
        conta_luz_futura = max(valor_conta_luz - economia_mensal, fatura_minima_tecnica)

    if conta_luz_futura > 0:
        conta_luz_futura = max(conta_luz_futura, fatura_minima_tecnica)

    if valor_conta_luz > 0 and valor_conta_luz < fatura_minima_tecnica:
        valor_conta_luz = fatura_minima_tecnica

    if conta_luz_futura > valor_conta_luz > 0:
        conta_luz_futura = valor_conta_luz

    economia_anual = _to_float_flex(getattr(projeto, "economia_anual", None), 0)
    if economia_anual <= 0 and economia_mensal > 0:
        economia_anual = economia_mensal * 12

    # Payback
    payback = payback_informado or (valor_total / economia_anual if economia_anual > 0 else 0)

    # ROI e economia em 25 anos
    economia_25 = economia_anual * 25
    roi_25 = ((economia_25 - valor_total) / valor_total) * 100 if valor_total > 0 else 0

    # Quantidade de módulos (tentar diferentes nomes de campo)
    qtd_modulos = (
        getattr(projeto, "qtd_placas", None)
        or getattr(projeto, "numero_paineis", None)
        or getattr(projeto, "qtd_modulos", None)
        or 0
    )

    # Área necessária (fallback: qtd_modulos * 2.5 m²)
    area = float(getattr(projeto, "area_necessaria", None) or 0)
    if area == 0 and qtd_modulos:
        area = float(qtd_modulos) * 2.5
    
    acrescimo_anual_percentual = _to_float_flex(
        getattr(projeto, 'acrescimo_anual_percentual', None)
        or getattr(projeto, 'reajuste_anual_energia', None)
        or getattr(projeto, 'reajuste_anual', None)
        or 10.0,
        10.0,
    )

    # Cliente
    cliente = None
    if MODELS_IMPORTED and hasattr(projeto, 'cliente_id') and projeto.cliente_id:
        try:
            cliente = Cliente.query.get(projeto.cliente_id)
        except Exception as e:
            logger.warning(f"Erro ao buscar cliente: {e}")
    
    if cliente:
        cliente_nome = (
            getattr(cliente, "nome_razao_social", None)
            or getattr(cliente, "razao_social", None)
            or getattr(cliente, "nome", None)
            or getattr(cliente, "nome_fantasia", None)
            or "Cliente não informado"
        )
    else:
        cliente_nome = "Cliente não informado"
    cliente_cidade = cliente.cidade if cliente else ""
    cliente_estado = cliente.estado if cliente else ""

    # Placa solar
    kit = None
    if MODELS_IMPORTED and hasattr(projeto, 'kit_id') and projeto.kit_id:
        try:
            kit = KitSolar.query.get(projeto.kit_id)
        except Exception as e:
            logger.warning(f"Erro ao buscar kit: {e}")

    placa = None
    if MODELS_IMPORTED and hasattr(projeto, 'placa_id') and projeto.placa_id:
        try:
            placa = PlacaSolar.query.get(projeto.placa_id)
        except Exception as e:
            logger.warning(f"Erro ao buscar placa: {e}")

    # Inversor
    inversor = None
    if MODELS_IMPORTED and hasattr(projeto, 'inversor_id') and projeto.inversor_id:
        try:
            inversor = InversorSolar.query.get(projeto.inversor_id)
        except Exception as e:
            logger.warning(f"Erro ao buscar inversor: {e}")

    if not placa and kit and getattr(kit, 'placa', None):
        placa = kit.placa
    if not inversor and kit and getattr(kit, 'inversor', None):
        inversor = kit.inversor

    if not qtd_modulos and kit and getattr(kit, 'qtd_placas', None):
        qtd_modulos = kit.qtd_placas

    modelo_modulo = (
        (placa.modelo if placa else None)
        or getattr(projeto, 'placa_modelo', None)
        or getattr(projeto, 'modulo_modelo', None)
        or getattr(projeto, 'modelo_modulo', None)
        or (kit.descricao if kit else None)
        or "Conforme kit selecionado"
    )
    fabricante_modulo = (
        (placa.fabricante if placa else None)
        or getattr(projeto, 'placa_fabricante', None)
        or getattr(projeto, 'fabricante_modulo', None)
        or (kit.fabricante if kit else None)
        or "-"
    )
    potencia_modulo_num = _extract_first_float(
        (placa.potencia if placa and getattr(placa, 'potencia', None) else None)
        or getattr(projeto, 'placa_potencia', None)
        or getattr(projeto, 'potencia_modulo', None),
        0,
    )
    if potencia_modulo_num <= 0 and kit and getattr(kit, 'potencia_kwp', None) and qtd_modulos:
        potencia_modulo_num = (_to_float_flex(kit.potencia_kwp, 0) * 1000.0) / max(int(qtd_modulos), 1)
    potencia_modulo = f"{potencia_modulo_num:.0f}W" if potencia_modulo_num > 0 else "-"

    modelo_inversor = (
        (inversor.modelo if inversor else None)
        or getattr(projeto, 'inversor_modelo', None)
        or "-"
    )
    fabricante_inversor = (
        (inversor.fabricante if inversor else None)
        or getattr(projeto, 'inversor_fabricante', None)
        or "-"
    )

    # Compatibilidade de campos do inversor entre versões/modelos
    inversor_potencia = "-"
    if inversor:
        pot_inv = (
            getattr(inversor, "potencia", None)
            or getattr(inversor, "potencia_nominal", None)
            or getattr(inversor, "potencia_maxima", None)
        )
        if pot_inv is not None:
            try:
                # Potência nominal do catálogo está em kW
                inversor_potencia = f"{formatar_numero(float(pot_inv), 2)} kW"
            except Exception:
                inversor_potencia = str(pot_inv)
    if inversor_potencia in ("", "-"):
        pot_inv = _extract_first_float(
            getattr(projeto, 'inversor_potencia', None)
            or getattr(projeto, 'potencia_inversor', None),
            0,
        )
        if pot_inv > 0:
            inversor_potencia = f"{formatar_numero(pot_inv, 2)} kW"

    garantia_inversor = "10 anos"
    if inversor and getattr(inversor, "garantia_anos", None):
        garantia_inversor = f"{int(inversor.garantia_anos)} anos"

    reducao_percentual = 0
    if valor_conta_luz > 0:
        reducao_percentual = ((valor_conta_luz - conta_luz_futura) / valor_conta_luz) * 100
    if reducao_percentual <= 0:
        reducao_percentual = 90

    kit_outras_inf = (
        (getattr(kit, 'outras_informacoes', None) if kit else None)
        or getattr(projeto, 'observacoes', None)
        or getattr(projeto, 'descricao', None)
        or ""
    )

    kit_desc = (
        (getattr(kit, 'descricao', None) if kit else None)
        or getattr(projeto, 'descricao', None)
        or ""
    )

    # Montar contexto completo com todos os placeholders
    contexto = {
        # Dados do cliente
        "NOME_CLIENTE": cliente_nome,
        "CPF_CNPJ_CLIENTE": cliente.cpf_cnpj if cliente and cliente.cpf_cnpj else "",
        "CIDADE": cliente_cidade or getattr(projeto, "localidade", "") or "",
        "ESTADO": cliente_estado or "",
        "ENDERECO_CLIENTE": cliente.endereco if cliente and cliente.endereco else "",
        "TELEFONE_CLIENTE": cliente.telefone if cliente and hasattr(cliente, 'telefone') else "",
        "EMAIL_CLIENTE": cliente.email if cliente and hasattr(cliente, 'email') else "",
        "CEP_CLIENTE": cliente.cep if cliente and hasattr(cliente, 'cep') else "",
        
        # Dados do projeto
        "NUMERO_PROJETO": str(getattr(projeto, "id", "")),
        "DATA_PROPOSTA": datetime.now().strftime("%d/%m/%Y"),
        "VALIDADE_PROPOSTA": (datetime.now() + timedelta(days=15)).strftime("%d/%m/%Y"),
        
        # Dados técnicos
        "POTENCIA_SISTEMA": f"{formatar_numero(getattr(projeto, 'potencia_kwp', 0), 2)} kWp",
        "QTD_MODULOS": str(qtd_modulos),
        "POTENCIA_MODULO": potencia_modulo,
        "MODELO_MODULO": modelo_modulo,
        "FABRICANTE_MODULO": fabricante_modulo,
        "MODELO_INVERSOR": modelo_inversor,
        "FABRICANTE_INVERSOR": fabricante_inversor,
        "POTENCIA_INVERSOR": inversor_potencia,
        "TIPO_INSTALACAO": getattr(projeto, 'tipo_instalacao', None) or "monofasica",
        "GARANTIA_MODULOS": "25 anos (potência) / 12 anos (produto)",
        "GARANTIA_INVERSOR": garantia_inversor,
        "VIDA_UTIL_SISTEMA": "25 anos",
        
        # Geração e consumo
        "GERACAO_MENSAL": f"{formatar_numero(getattr(projeto, 'geracao_estimada_mes', 0), 0)} kWh/mês",
        "GERACAO_ANUAL": f"{formatar_numero(float(getattr(projeto, 'geracao_estimada_mes', 0) or 0) * 12, 0)} kWh/ano",
        "CONSUMO_MENSAL": f"{formatar_numero(consumo_mensal, 0)} kWh/mês",
        "CONSUMO_ANUAL": f"{formatar_numero(consumo_mensal * 12, 0)} kWh/ano",
        "AREA_NECESSARIA": f"{formatar_numero(area, 2)} m²",
        "IRRADIACAO_SOLAR": f"{formatar_numero(getattr(projeto, 'irradiacao_solar', 5.0), 2)} kWh/m².dia",
        "PRECO_KWH": f"R$ {formatar_numero(preco_kwh, 4)}",
        "TARIFA_ENERGIA": f"R$ {formatar_numero(preco_kwh, 4)}",
        "tarifa_kwh": f"{preco_kwh:.4f}",
        "CONSUMO_MINIMO_KWH": f"{consumo_minimo_kwh:.0f}",
        "consumo_minimo_kwh": f"{consumo_minimo_kwh:.0f}",
        "FATURA_MINIMA_TECNICA": formatar_moeda(fatura_minima_tecnica),
        "fatura_minima_tecnica": f"{fatura_minima_tecnica:.2f}",
        "ADICIONAIS_FATURA": formatar_moeda(adicionais_fatura),
        "adicionais_fatura": f"{adicionais_fatura:.2f}",
        
        # Valores financeiros
        "VALOR_INVESTIMENTO": formatar_moeda(valor_total),
        "ECONOMIA_MENSAL": formatar_moeda(economia_mensal),
        "ECONOMIA_ANUAL": formatar_moeda(economia_anual),
        "ECONOMIA_25_ANOS": formatar_moeda(economia_25),
        "PAYBACK": f"{formatar_numero(payback, 1)} anos",
        "ROI_25_ANOS": f"{formatar_numero(roi_25, 0)}%",
        "CONTA_LUZ_ATUAL": formatar_moeda(valor_conta_luz),
        "CONTA_LUZ_FUTURA": formatar_moeda(conta_luz_futura),
        "REDUCAO_PERCENTUAL": f"{formatar_numero(reducao_percentual, 0)}%",
        "PERCENTUAL_COMPENSACAO": formatar_numero(float(getattr(projeto, 'simultaneity_factor', None) or 100), 0) + "%",
        "FATURA_SEM_SISTEMA": formatar_moeda(valor_conta_luz),
        "FATURA_COM_SISTEMA": formatar_moeda(conta_luz_futura),
        "fatura_sem_sistema": f"{valor_conta_luz:.2f}",
        "fatura_com_sistema": f"{conta_luz_futura:.2f}",
        "fatura_minima": f"{conta_luz_futura:.2f}",
        "fatura_media_mensal_sem_sistema": f"{valor_conta_luz:.2f}",
        "fatura_sem_sistema_rs": formatar_moeda(valor_conta_luz),
        "fatura_com_sistema_rs": formatar_moeda(conta_luz_futura),
        "ACRESCIMO_ANUAL_PERCENTUAL": f"{formatar_numero(acrescimo_anual_percentual, 2)}%",
        "acrescimo_anual_percentual": f"{acrescimo_anual_percentual:.2f}",
        "REAJUSTE_ANUAL": f"{formatar_numero(acrescimo_anual_percentual, 2)}%",
        "reajuste_anual": f"{acrescimo_anual_percentual:.2f}",

        # Campos adicionais de kit/outros detalhes
        "KIT_DESCRICAO": kit_desc,
        "kit_descricao": kit_desc,
        "kit_outras_inf": kit_outras_inf,
        "kit_outras_informacoes": kit_outras_inf,
        "outras_informacoes_kit": kit_outras_inf,
        "outras_informacoes": kit_outras_inf,
        "OUTRAS_DESCRICOES": kit_outras_inf,
        "outras_descricoes": kit_outras_inf,
        
        # Custos detalhados
        "CUSTO_EQUIPAMENTOS": formatar_moeda(getattr(projeto, 'custo_equipamentos', 0)),
        "CUSTO_INSTALACAO": formatar_moeda(getattr(projeto, 'custo_instalacao', 0)),
        "CUSTO_PROJETO": formatar_moeda(getattr(projeto, 'custo_projeto', 0)),
        
        # Condições comerciais
        "FORMA_PAGAMENTO": getattr(projeto, 'forma_pagamento', None) or "À vista ou parcelado",
        "PRAZO_ENTREGA": getattr(projeto, 'prazo_entrega', None) or "45 dias úteis",
        "PRAZO_INSTALACAO": "7 a 15 dias úteis após entrega dos equipamentos",
        "PRAZO_TOTAL": "60 dias úteis (aproximadamente)",
        
        # Dados da empresa (configuração)
        "NOME_EMPRESA": "JSP Elétrica & Solar",
        "CNPJ_EMPRESA": "",
        "TELEFONE_EMPRESA": "(15) 99670-2036",
        "EMAIL_EMPRESA": "atendimento@eletricasaroba.com.br",
        "SITE_EMPRESA": "",
    }

    # Tentar carregar dados da empresa de configuração
    if CONFIG_IMPORTED:
        try:
            config = carregar_configuracao()
            if config:
                contexto["NOME_EMPRESA"] = config.nome_fantasia or contexto["NOME_EMPRESA"]
                contexto["CNPJ_EMPRESA"] = config.cnpj or ""
                contexto["TELEFONE_EMPRESA"] = config.telefone or contexto["TELEFONE_EMPRESA"]
                contexto["EMAIL_EMPRESA"] = config.email or contexto["EMAIL_EMPRESA"]
                contexto["SITE_EMPRESA"] = config.site or ""
        except Exception as e:
            logger.warning(f"Não foi possível carregar configurações da empresa: {e}")

    return contexto


def gerar_docx_proposta(projeto, template_path, output_path):
    """
    Carrega template Word, substitui placeholders e salva documento preenchido.
    
    Args:
        projeto: Objeto ProjetoSolar com dados do projeto
        template_path: Caminho para o arquivo .docx modelo
        output_path: Caminho onde salvar o .docx preenchido
        
    Returns:
        Path do arquivo gerado
    """
    logger.info(f"🔧 Montando contexto do projeto {projeto.id}")
    contexto = montar_contexto_proposta(projeto)

    # Usa o mesmo motor robusto da rotina de upload para suportar:
    # - múltiplos formatos de placeholder
    # - fallback XML
    # - placeholders visuais de gráficos/tabelas
    from app.energia_solar.word_utils import substituir_variaveis_word

    logger.info(f"📄 Carregando template Word: {template_path}")
    logger.info(f"✏️ Substituindo {len(contexto)} placeholders no documento")
    doc = substituir_variaveis_word(template_path, contexto)
    
    logger.info(f"💾 Salvando documento preenchido: {output_path}")
    doc.save(output_path)

    # Fallback para elementos não cobertos pelo python-docx (ex.: text boxes)
    arquivos_xml_alterados = substituir_placeholders_xml_docx(output_path, contexto)
    if arquivos_xml_alterados > 0:
        logger.info(f"✅ Fallback XML aplicado em {arquivos_xml_alterados} arquivo(s) interno(s) do DOCX")
    
    logger.info(f"✅ Documento DOCX gerado com sucesso!")
    return output_path


def converter_docx_para_pdf(docx_path, pdf_path):
    """
    Converte DOCX para PDF usando LibreOffice (se disponível).
    
    Args:
        docx_path: Caminho do arquivo .docx
        pdf_path: Caminho desejado para o .pdf
        
    Returns:
        Path do PDF gerado ou None se falhar
    """
    try:
        output_dir = str(Path(pdf_path).parent)
        
        logger.info(f"📝 Convertendo DOCX para PDF usando LibreOffice...")
        
        # Tentar usar LibreOffice para conversão
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            str(docx_path)
        ], check=True, timeout=60, capture_output=True)
        
        # LibreOffice gera o PDF com o mesmo nome do DOCX
        gerado = Path(output_dir) / (Path(docx_path).stem + ".pdf")
        
        if gerado.exists():
            # Renomear se necessário
            if str(gerado) != str(pdf_path):
                gerado.rename(pdf_path)
            
            logger.info(f"✅ PDF gerado com sucesso: {pdf_path}")
            return pdf_path
        else:
            logger.warning(f"⚠️ LibreOffice executou mas PDF não foi encontrado")
            return None

    except FileNotFoundError:
        logger.warning(f"⚠️ LibreOffice não encontrado no sistema. PDF não será gerado.")
        return None
    except subprocess.TimeoutExpired:
        logger.error(f"❌ Timeout ao converter DOCX para PDF (60s)")
        return None
    except Exception as e:
        logger.error(f"❌ Erro ao converter DOCX para PDF: {e}")
        return None
