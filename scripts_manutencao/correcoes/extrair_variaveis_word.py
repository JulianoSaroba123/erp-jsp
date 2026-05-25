"""
Script para extrair todas as variáveis de um documento Word
Identifica padrões como: {{VARIAVEL}}, [VARIAVEL], {VARIAVEL}, <VARIAVEL>
"""

from docx import Document
import re
from pathlib import Path

def extrair_variaveis_documento(caminho_docx):
    """Extrai todas as variáveis encontradas no documento Word"""
    
    print(f"📄 Abrindo documento: {caminho_docx}")
    doc = Document(caminho_docx)
    
    # Padrões para identificar variáveis
    padroes = [
        (r'\{\{([A-Z_0-9]+)\}\}', '{{VARIAVEL}}'),  # {{VARIAVEL}}
        (r'\[([A-Z_][A-Z_0-9]*)\]', '[VARIAVEL]'),  # [VARIAVEL]
        (r'\{([A-Z_][A-Z_0-9]*)\}', '{VARIAVEL}'),  # {VARIAVEL}
        (r'<([A-Z_][A-Z_0-9]*)>', '<VARIAVEL>'),    # <VARIAVEL>
        (r'\$\{([A-Z_][A-Z_0-9]*)\}', '${VARIAVEL}'),  # ${VARIAVEL}
        (r'\%([A-Z_][A-Z_0-9]*)\%', '%VARIAVEL%'),  # %VARIAVEL%
    ]
    
    variaveis_encontradas = {}  # dict para guardar variável e seu formato
    
    # Buscar em parágrafos
    print("\n🔍 Buscando variáveis nos parágrafos...")
    for i, paragrafo in enumerate(doc.paragraphs, 1):
        texto = paragrafo.text
        if texto.strip():
            for padrao, formato in padroes:
                matches = re.findall(padrao, texto)
                if matches:
                    print(f"  Parágrafo {i}: {matches} (formato: {formato})")
                    for match in matches:
                        if match not in variaveis_encontradas:
                            variaveis_encontradas[match] = formato
    
    # Buscar em tabelas
    print("\n🔍 Buscando variáveis nas tabelas...")
    for i_tabela, tabela in enumerate(doc.tables, 1):
        for i_linha, linha in enumerate(tabela.rows, 1):
            for i_celula, celula in enumerate(linha.cells, 1):
                texto = celula.text
                if texto.strip():
                    for padrao, formato in padroes:
                        matches = re.findall(padrao, texto)
                        if matches:
                            print(f"  Tabela {i_tabela}, Linha {i_linha}, Célula {i_celula}: {matches} (formato: {formato})")
                            for match in matches:
                                if match not in variaveis_encontradas:
                                    variaveis_encontradas[match] = formato
    
    # Buscar em cabeçalhos e rodapés
    print("\n🔍 Buscando variáveis nos cabeçalhos e rodapés...")
    for secao in doc.sections:
        # Cabeçalho
        for paragrafo in secao.header.paragraphs:
            texto = paragrafo.text
            if texto.strip():
                for padrao, formato in padroes:
                    matches = re.findall(padrao, texto)
                    if matches:
                        print(f"  Cabeçalho: {matches} (formato: {formato})")
                        for match in matches:
                            if match not in variaveis_encontradas:
                                variaveis_encontradas[match] = formato
        
        # Rodapé
        for paragrafo in secao.footer.paragraphs:
            texto = paragrafo.text
            if texto.strip():
                for padrao, formato in padroes:
                    matches = re.findall(padrao, texto)
                    if matches:
                        print(f"  Rodapé: {matches} (formato: {formato})")
                        for match in matches:
                            if match not in variaveis_encontradas:
                                variaveis_encontradas[match] = formato
    
    return variaveis_encontradas


if __name__ == "__main__":
    # Caminho do documento
    caminho = r"E:\Modelos de Documentos_Solar\Documento-JSPELE-0002 - Propsta Modelo 02-Usual.docx"
    
    try:
        variaveis_dict = extrair_variaveis_documento(caminho)
        variaveis = sorted(list(variaveis_dict.keys()))
        
        print("\n" + "="*80)
        print(f"✅ TOTAL DE VARIÁVEIS ENCONTRADAS: {len(variaveis)}")
        print("="*80)
        
        if variaveis:
            print("\n📋 LISTA DE VARIÁVEIS COM SEUS FORMATOS:\n")
            for i, var in enumerate(variaveis, 1):
                formato = variaveis_dict[var]
                print(f"{i:3d}. {var:30s} → {formato}")
            
            print("\n" + "="*80)
            print("📝 FORMATO {{{{VARIAVEL}}}} PARA COPIAR NO WORD:")
            print("="*80)
            for var in variaveis:
                print("{{" + var + "}}")
                
        else:
            print("\n⚠️ Nenhuma variável encontrada no documento!")
            print("Padrões buscados: {{VARIAVEL}}, [VARIAVEL], {VARIAVEL}, <VARIAVEL>, ${VARIAVEL}, %VARIAVEL%")
    
    except FileNotFoundError:
        print(f"\n❌ ERRO: Arquivo não encontrado!")
        print(f"Caminho: {caminho}")
        print("\nVerifique se:")
        print("1. O arquivo existe no local especificado")
        print("2. O caminho está correto (use raw string ou barras duplas)")
    
    except Exception as e:
        print(f"\n❌ ERRO ao processar documento: {e}")
        import traceback
        traceback.print_exc()
